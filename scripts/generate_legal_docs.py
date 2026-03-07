"""
Generates production-ready Ocypheris legal DOCX files.
Run: python3 scripts/generate_legal_docs.py
Output: docs/legal/Ocypheris_Privacy_Policy.docx
        docs/legal/Ocypheris_Terms_and_Conditions.docx
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), '..', 'docs', 'legal')
os.makedirs(OUTPUT_DIR, exist_ok=True)


def set_heading_style(para, level=1):
    para.style = f'Heading {level}'
    run = para.runs[0] if para.runs else para.add_run()
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)  # dark slate
    if level == 1:
        run.font.size = Pt(20)
    elif level == 2:
        run.font.size = Pt(14)
        run.bold = True
    else:
        run.font.size = Pt(12)
        run.bold = True
        run.italic = True


def add_cover(doc, title, effective_date):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('OCYPHERIS')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    doc.add_paragraph()

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(title)
    run2.font.size = Pt(22)
    run2.font.bold = True
    run2.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)

    doc.add_paragraph()

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(f'Effective Date: {effective_date}')
    run3.font.size = Pt(11)
    run3.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run('Autopilot — AWS Security Operations Platform')
    run4.font.size = Pt(11)
    run4.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    run4.italic = True

    doc.add_page_break()


def add_section(doc, number, title, paragraphs_or_content):
    h = doc.add_heading(f'{number}. {title}', level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
        run.font.size = Pt(14)

    if isinstance(paragraphs_or_content, str):
        doc.add_paragraph(paragraphs_or_content)
    elif isinstance(paragraphs_or_content, list):
        for item in paragraphs_or_content:
            if isinstance(item, tuple) and item[0] == 'sub':
                h2 = doc.add_heading(item[1], level=2)
                for run in h2.runs:
                    run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
                    run.font.size = Pt(12)
                for line in item[2:]:
                    if isinstance(line, list):
                        for bullet in line:
                            p = doc.add_paragraph(style='List Bullet')
                            p.add_run(bullet)
                    else:
                        doc.add_paragraph(line)
            elif isinstance(item, list):
                for bullet in item:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(bullet)
            elif isinstance(item, str):
                doc.add_paragraph(item)

    doc.add_paragraph()


###############################################################################
# PRIVACY POLICY
###############################################################################

def build_privacy_policy():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Default body font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    add_cover(doc, 'Privacy Policy', 'December 13, 2025')

    add_section(doc, 1, 'Introduction', [
        'Ocypheris ("Ocypheris," "we," "us," or "our") is a software-as-a-service company offering the Autopilot platform — an automated AWS security operations platform that ingests findings from AWS Security Hub, Amazon GuardDuty, and IAM Access Analyzer, converts them into prioritised remediation actions, and allows authorised users to execute approved fixes directly against connected AWS environments.',
        'This Privacy Policy describes how we collect, use, store, share, and protect personal and technical data when you access or use our platform at ocypheris.com and any related APIs or services (collectively, the "Services"). By using the Services you agree to the practices described here. If you do not agree, do not access or use the Services.',
    ])

    add_section(doc, 2, 'Data We Collect', [
        ('sub', '2.1  Account & Identity Data',
         'When you register for Autopilot, we collect:',
         ['Full name, work email address, and password hash (never stored in plain text)',
          'Company or organisation name',
          'IP address and approximate geographic location at time of registration and login',
          'User role and invite metadata (for team management)']),
        ('sub', '2.2  AWS Account Integration Data',
         'To connect your AWS environment to Autopilot, you create and provide an IAM role that Ocypheris assumes via AWS Security Token Service (STS). We collect and store:',
         ['Your AWS Account ID(s)',
          'IAM Role ARNs for the Read Role and Write Role',
          'The ExternalId value associated with each role (used as a per-tenant security token)',
          'AWS region(s) you enable for scanning',
          'Connection status, last-ingest timestamps, and control-plane readiness metadata'],
         'We do not store long-lived AWS access keys or secret access keys. All access to your AWS environment uses short-lived STS-issued credentials, rotated automatically per session.'),
        ('sub', '2.3  Security Findings & Actions Data',
         'When Autopilot ingests your AWS environment, we receive and store data generated by AWS security services within your account:',
         ['Security Hub findings (all supported controls and standards)',
          'GuardDuty threat-detection findings',
          'IAM Access Analyzer findings',
          'Normalised action records derived from findings (deduplication, priority scores, grouping)',
          'Exception records, including the approving user, reason, and expiry date'],
         'This data is derived from your AWS environment and constitutes your organisation\'s security posture information. It is treated as confidential business data.'),
        ('sub', '2.4  Remediation Run Data',
         'When you approve and execute a remediation action, we log:',
         ['The action type, target AWS resource, region, and account',
          'The approving user and timestamp',
          'Execution mode (direct fix or IaC patch bundle)',
          'Pre-check and post-check results',
          'Full execution log and outcome status',
          'Artefact URLs for generated IaC patch bundles (stored in Amazon S3 within our AWS account)']),
        ('sub', '2.5  Evidence & Compliance Export Data',
         'When you generate a compliance or evidence pack, we create and store export artefacts including finding snapshots, exception attestations, control mappings, and auditor summaries. These are written to Amazon S3 in our AWS account and a download link is made available to you.'),
        ('sub', '2.6  Usage & Platform Telemetry',
         'We collect operational data about your use of the platform, including:',
         ['Pages visited, features used, and actions taken within the Autopilot dashboard',
          'API request logs, error rates, and response latencies',
          'Browser type, operating system, and device category',
          'Session identifiers (stored in browser memory/localStorage; not set as persistent third-party cookies)']),
        ('sub', '2.7  Communication & Support Data',
         ['Emails you send to us, including support requests',
          'Digest email delivery logs (weekly security digest, sent to addresses you configure)',
          'Slack webhook configuration (if you enable Slack digest notifications)']),
        ('sub', '2.8  Billing & Payment Data',
         'We process subscription payments via bank transfer. We collect and store invoice records, payment status, and the billing contact you designate. We do not store banking credentials or payment account numbers.'),
    ])

    add_section(doc, 3, 'How We Use Your Data', [
        'We use the data we collect to:',
        ['Deliver the Services — authenticate users, connect to your AWS environment, ingest findings, run action engine jobs, execute approved remediations, and generate evidence packs',
         'Operate and improve the platform — monitor performance, diagnose errors, develop new features, and train internal models on aggregate and anonymised metrics (never on your raw findings without consent)',
         'Communicate with you — send weekly security digests, Slack notifications, product updates, billing notices, and security alerts to email addresses you configure',
         'Fulfil legal and compliance obligations — maintain audit trails, respond to lawful requests, and comply with applicable laws',
         'Protect security and prevent misuse — detect and respond to fraudulent or abusive use of the platform'],
        'Autopilot uses automated processing to score, prioritise, and surface remediation recommendations. These automated scores are advisory tools to support human decision-making. No remediation action is executed without explicit human approval within the platform.',
    ])

    add_section(doc, 4, 'Legal Bases for Processing (GDPR)', [
        'For users in the European Economic Area (EEA) and United Kingdom, we process personal data under the following legal bases:',
        ['Contract performance — processing necessary to deliver the Services you subscribed to',
         'Legitimate interests — platform security, fraud prevention, and product analytics, provided these are not overridden by your rights',
         'Legal obligation — where required by applicable law',
         'Consent — for optional communications such as marketing emails, where we will request separate consent'],
    ])

    add_section(doc, 5, 'Data Sharing', [
        'We do not sell your personal data. We share data only in the following circumstances:',
        ('sub', '5.1  AWS Services (on your behalf)',
         'When Autopilot accesses your AWS environment, it interacts with AWS APIs (Security Hub, GuardDuty, IAM, STS, S3, etc.) using credentials derived from the role you grant. This access occurs within your AWS account under your account policies.'),
        ('sub', '5.2  Infrastructure & Hosting Providers',
         'Our platform is hosted on Amazon Web Services in the eu-north-1 (Stockholm) region. AWS acts as a sub-processor for data stored on our platform. All data is encrypted at rest and in transit.'),
        ('sub', '5.3  Email Delivery',
         'We use a transactional email provider to deliver system notifications and weekly digest emails. The provider receives email addresses and message content where necessary to deliver messages on our behalf.'),
        ('sub', '5.4  Legal & Regulatory Requirements',
         'We may disclose your data to law enforcement, regulators, or other third parties if required by Egyptian law, EU law (where applicable), or other applicable legal authority, or to protect the legal rights and safety of Ocypheris, its users, or the public.'),
        ('sub', '5.5  Business Transfers',
         'If Ocypheris undergoes a merger, acquisition, or sale of assets, your data may be transferred as part of that transaction. We will notify you before your data is subject to a different privacy policy.'),
    ])

    add_section(doc, 6, 'Data Retention', [
        ['Account data: retained for the duration of your subscription and for 90 days after account closure',
         'Security findings & actions: retained for the duration of your subscription; deleted within 60 days of account closure unless you request an earlier deletion',
         'Remediation run logs & evidence packs: retained for the duration of your subscription plus 12 months (to support post-audit review); deleted on your written request after account closure',
         'Billing records: retained for 7 years as required by applicable financial record-keeping obligations',
         'Backups: encrypted database snapshots are retained for up to 30 days and automatically purged'],
        'You may request deletion of your data at any time by contacting legal@ocypheris.com. We will fulfil deletion requests within 30 days except where retention is required by law.',
    ])

    add_section(doc, 7, 'Data Security', [
        'We implement security controls appropriate to the sensitivity of the data we process, including:',
        ['TLS 1.2+ encryption for all data in transit between your browser, our API, and AWS services',
         'AES-256 encryption for all data at rest (RDS Postgres, S3)',
         'Short-lived STS credentials for all AWS access — no long-lived AWS keys are stored',
         'Strict tenant isolation: row-level access controls ensure each organisation can only access its own data',
         'Least-privilege IAM policies for Autopilot\'s own AWS service access',
         'CloudWatch monitoring, alerting, and audit logging for all platform activity',
         'AWS Secrets Manager for storage of all platform secrets and credentials'],
        'No security system is perfect. In the event of a data breach that affects your personal data, we will notify you and the relevant supervisory authority within 72 hours of discovery, where required by law.',
    ])

    add_section(doc, 8, 'International Data Transfers', [
        'Ocypheris is incorporated in Egypt. Your data is processed and stored on AWS infrastructure in the eu-north-1 (Stockholm, Sweden) region, within the European Union. If your use of the platform involves a transfer of personal data from the EEA or UK to a country not deemed adequate by the European Commission or the UK Government (including Egypt), we rely on Standard Contractual Clauses (SCCs) and the UK International Data Transfer Addendum as our legal transfer mechanisms. You may request a copy of these by contacting legal@ocypheris.com.',
    ])

    add_section(doc, 9, 'Your Rights', [
        ('sub', '9.1  All Users',
         ['Access: Request a copy of the personal data we hold about you',
          'Correction: Request correction of inaccurate or incomplete data',
          'Deletion: Request deletion of your personal data, subject to legal retention obligations',
          'Data portability: Receive your data in a structured, machine-readable format (JSON/CSV)']),
        ('sub', '9.2  EEA & UK Users (GDPR / UK GDPR)',
         ['Restriction of processing: Request that we limit how we process your data in certain circumstances',
          'Object to processing: Object to processing based on legitimate interests',
          'Withdraw consent: Where processing is based on consent, withdraw it at any time without affecting prior processing',
          'Lodge a complaint: With your national data protection authority (DPA)']),
        ('sub', '9.3  California Users (CCPA / CPRA)',
         ['Right to know what personal information is collected, used, shared, or sold',
          'Right to delete your personal information',
          'Right to opt-out of the sale of personal information (we do not sell your data)',
          'Right to non-discrimination for exercising your rights']),
        'To exercise any of these rights, contact legal@ocypheris.com. We will respond within 30 days.',
    ])

    add_section(doc, 10, 'Cookies & Tracking', [
        'Autopilot uses minimal browser storage to operate:',
        ['Authentication token: A JWT token stored in browser localStorage to keep you signed in. This is a functional necessity and not a tracking cookie.',
         'Session preferences: Theme and UI preference settings stored in localStorage.'],
        'We do not use third-party advertising cookies or sell browsing data. If we introduce analytics cookies in the future, we will update this policy and seek consent where required by law. You can clear localStorage at any time through your browser settings; this will sign you out of the platform.',
    ])

    add_section(doc, 11, "Children's Privacy", [
        'The Services are intended for business users aged 18 and over. We do not knowingly collect personal information from individuals under 18. If you believe we have inadvertently collected such data, please contact legal@ocypheris.com and we will delete it promptly.',
    ])

    add_section(doc, 12, 'Changes to This Policy', [
        'We may update this Privacy Policy from time to time to reflect changes in the Services, applicable law, or our data practices. We will notify you of material changes by email to your registered address and by posting a prominent notice on ocypheris.com at least 14 days before the change takes effect. Your continued use of the Services after the effective date constitutes acceptance of the updated policy. If you do not agree to the updated policy, you must discontinue use and contact us to close your account.',
    ])

    add_section(doc, 13, 'Contact Us', [
        'If you have any questions, concerns, or requests regarding this Privacy Policy or your personal data, please contact:',
        'Ocypheris — Data Privacy Team\nEmail: legal@ocypheris.com\nWebsite: ocypheris.com',
        'For EEA users, if you believe your rights have not been respected, you also have the right to file a complaint with the data protection authority in your country of residence.',
    ])

    out_path = os.path.join(OUTPUT_DIR, 'Ocypheris_Privacy_Policy.docx')
    doc.save(out_path)
    print(f'Saved: {out_path}')


###############################################################################
# TERMS & CONDITIONS
###############################################################################

def build_terms():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    add_cover(doc, 'Terms of Service', 'December 13, 2025')

    add_section(doc, 1, 'Introduction & Acceptance', [
        'These Terms of Service ("Terms") constitute a legally binding agreement between you (the subscriber, "Customer," "you," or "your") and Ocypheris ("Ocypheris," "we," "us," or "our") governing your access to and use of the Autopilot platform available at ocypheris.com and any related APIs, services, workers, or exported artefacts (collectively, the "Services").',
        'By creating an account, clicking "Sign Up," connecting an AWS account, or otherwise accessing the Services, you acknowledge that you have read, understood, and agree to be bound by these Terms and our Privacy Policy (available at ocypheris.com/legal/privacy), which is incorporated herein by reference. If you are entering into these Terms on behalf of an organisation, you represent that you have authority to bind that organisation. If you do not agree to these Terms, do not access or use the Services.',
    ])

    add_section(doc, 2, 'Definitions', [
        ['"Autopilot" means the Ocypheris cloud security SaaS platform that ingests AWS security findings, surfaces prioritised remediation actions, and executes approved remediations.',
         '"AWS Environment" means the Amazon Web Services account(s), infrastructure, configurations, and resources owned or operated by the Customer that are connected to Autopilot.',
         '"Read Role" means an AWS IAM role with read-only permissions, deployed by the Customer in their AWS account, which Autopilot uses to ingest security findings.',
         '"Write Role" means an AWS IAM role with limited write permissions, deployed by the Customer in their AWS account, which Autopilot uses solely to execute direct-fix remediations approved by the Customer.',
         '"Finding" means a security vulnerability, misconfiguration, or risk signal ingested from AWS Security Hub, Amazon GuardDuty, or IAM Access Analyzer.',
         '"Action" means a prioritised, deduplicated task derived from one or more Findings and presented to the Customer for review, exception, or remediation.',
         '"Remediation Run" means an execution event — initiated by an explicit Customer approval — in which Autopilot applies a fix to the Customer\'s AWS Environment via the Write Role, or generates an IaC patch bundle for Customer-managed deployment.',
         '"IaC Patch Bundle" means a Terraform or CloudFormation template or patch generated by Autopilot and delivered to the Customer for independent review and deployment.',
         '"Evidence Pack" means a compliance export artefact (findings, exceptions, control mappings, attestations) generated by Autopilot and stored in Amazon S3.',
         '"Subscription" means the paid plan tier governing the Customer\'s access to the Services as agreed in the applicable order form or subscription agreement.'],
    ])

    add_section(doc, 3, 'The Services', [
        'Ocypheris provides the Autopilot platform as a subscription-based SaaS service. The Services include:',
        ['AWS account onboarding and IAM role-based integration (Read Role + Write Role)',
         'Continuous ingestion of security findings from AWS Security Hub, GuardDuty, and IAM Access Analyzer',
         'Automated action grouping, deduplication, and context-driven prioritisation',
         'Exception management with configurable expiry and approval workflows',
         'Hybrid remediation: direct-fix execution (via Write Role, with Customer approval) and IaC patch bundle generation',
         'Compliance and evidence pack export (CSV, JSON, ZIP)',
         'Weekly security digest notifications via email and/or Slack',
         'Baseline security reports'],
        'Ocypheris reserves the right to modify, enhance, or discontinue features of the Services with reasonable notice to active Customers. We will not materially reduce the core functionality of your subscribed plan without offering a refund for the affected period.',
    ])

    add_section(doc, 4, 'AWS Access & Customer Authorisation', [
        ('sub', '4.1  Granting Access',
         'To use Autopilot, you must deploy a Read Role and a Write Role in your AWS account(s), configured with an ExternalId provided by Ocypheris. By deploying these roles and connecting your account through the Autopilot platform, you explicitly authorise Ocypheris to:',
         ['Assume the Read Role to ingest and analyse security findings from your AWS Environment on an ongoing, automated basis',
          'Assume the Write Role solely to execute Remediation Runs you have individually approved within the platform']),
        ('sub', '4.2  No Unauthorised Actions',
         'Ocypheris will not use your Write Role to execute any action that has not been explicitly approved by an authorised user in your Autopilot account. We will not modify, delete, or reconfigure any AWS resource beyond the specific action scope of an approved Remediation Run. Each Remediation Run is logged with full before-and-after state, execution output, and the identity of the approving user.'),
        ('sub', '4.3  Credential Security',
         'Ocypheris accesses your AWS Environment exclusively via STS AssumeRole with ExternalId. No long-lived AWS access keys or secret keys are stored by Ocypheris. All temporary credentials are discarded after each job execution. You may revoke access at any time by modifying or deleting the IAM roles in your AWS account.'),
        ('sub', '4.4  Customer Representation',
         'By connecting an AWS account you represent and warrant that:',
         ['You are the account owner or have been authorised by the account owner to grant this access',
          'You have the legal and organisational authority to permit automated security operations on those AWS resources',
          'The AWS environment you connect is not owned by or shared with any party whose consent you have not obtained']),
    ])

    add_section(doc, 5, 'Automated Remediation — Customer Responsibility', [
        'IMPORTANT NOTICE: Autopilot has the ability to make changes to your live AWS environment when you approve a Remediation Run. Please read this section carefully before approving any remediation.',
        ('sub', '5.1  Approval-Gated Execution',
         'Every Remediation Run requires explicit approval by an authorised user within your Autopilot account. No remediation is executed automatically or without your consent. By approving a Remediation Run, you instruct Ocypheris to execute the described action against your AWS Environment and you assume responsibility for the consequences of that instruction.'),
        ('sub', '5.2  Pre-Checks and Logging',
         'Prior to executing a direct-fix remediation, Autopilot performs configurable safety pre-checks and logs the original state of the affected resource. Post-execution verification checks are run where applicable, and a full execution log is retained.'),
        ('sub', '5.3  Customer Backup Responsibility',
         'YOU ARE SOLELY RESPONSIBLE for maintaining independent backups, snapshots, and disaster recovery plans for your AWS Environment, independent of any Autopilot engagement. Ocypheris does not create backups of your AWS resources before executing remediations. You should verify that your environment has appropriate recovery mechanisms before approving any Remediation Run.'),
        ('sub', '5.4  No Guarantee of Remediation Outcome',
         'While Autopilot implements safety checks and follows AWS best practices, no automated system can guarantee that a remediation action will produce the exact desired outcome in every possible AWS configuration. AWS service behaviour, race conditions, regional dependencies, or custom configurations may produce unexpected results. You should review rollback guidance provided with each action type before approving execution.'),
        ('sub', '5.5  IaC Patch Bundles',
         'IaC Patch Bundles are generated templates for Customer-managed deployment. Ocypheris provides these outputs as reference artefacts. The Customer bears sole responsibility for reviewing, testing, and deploying IaC Patch Bundles in their environment.'),
    ])

    add_section(doc, 6, 'Customer Responsibilities', [
        ['Provide accurate, current, and complete information when registering and throughout the term of your Subscription',
         'Maintain the security of your Autopilot account credentials and promptly notify us of any suspected unauthorised access',
         'Ensure that all users accessing your Autopilot account have agreed to these Terms',
         'Use the Services only for lawful purposes in compliance with all applicable laws and regulations',
         'Connect only AWS accounts that you own or have explicit authority to manage',
         'Review and make independent decisions regarding all Findings, Actions, and Remediation Runs presented by Autopilot',
         'Maintain current contact information so we can notify you of material changes, security events, or service notices'],
    ])

    add_section(doc, 7, 'Acceptable Use', [
        'You must not use the Services to:',
        ['Connect or scan AWS accounts you do not own or have written, explicit authorisation to manage',
         'Reverse engineer, decompile, or attempt to extract the source code of any part of the platform',
         'Attempt to circumvent tenant isolation, access other customers\' data, or exploit platform vulnerabilities',
         'Use the platform for any purpose that violates applicable law, including data protection, export control, or sanctions law',
         'Resell, sublicense, or provide access to the Services to any third party without Ocypheris\' prior written consent',
         'Use automated mechanisms (bots, scrapers) to access the platform in a manner that degrades service for other users'],
        'Ocypheris reserves the right to suspend or terminate access immediately and without notice if we reasonably believe you are in violation of this section.',
    ])

    add_section(doc, 8, 'Intellectual Property', [
        ('sub', '8.1  Ocypheris IP',
         'The Autopilot platform, including its software, algorithms, action scoring methodology, user interface, documentation, and all underlying technology, is the exclusive intellectual property of Ocypheris and is protected by applicable copyright, trade secret, and intellectual property laws. Nothing in these Terms transfers any ownership of Ocypheris IP to you. You receive a limited, non-exclusive, non-transferable, revocable licence to access and use the Services during your active Subscription solely for your internal business security operations.'),
        ('sub', '8.2  Customer Data',
         'You retain all ownership of your AWS Environment data, security findings, and any proprietary business information you bring to the platform. You grant Ocypheris a limited licence to process, store, and use this data solely to provide and improve the Services as described in our Privacy Policy.'),
        ('sub', '8.3  Feedback',
         'If you provide suggestions, feedback, or ideas regarding the Services, you grant Ocypheris a perpetual, irrevocable, royalty-free licence to use such feedback for any purpose without restriction or compensation.'),
    ])

    add_section(doc, 9, 'Subscription, Fees & Payment', [
        ('sub', '9.1  Subscription Plans',
         'Access to Autopilot is offered on a subscription basis. Plan tiers, features, and pricing are specified in your applicable order form or subscription agreement entered into with Ocypheris.'),
        ('sub', '9.2  Payment',
         'Fees are payable via bank transfer as specified on your invoice. Unless otherwise agreed in writing, invoices are due within 30 days of the invoice date. All fees are stated exclusive of applicable taxes. You are responsible for all taxes, duties, or levies imposed on the Services by applicable authority.'),
        ('sub', '9.3  Late Payment',
         'Unpaid invoices after the due date may accrue interest at 1.5% per month (or the maximum permitted by applicable law, whichever is lower). Ocypheris reserves the right to suspend access to the Services for accounts with overdue balances after 15 days\' written notice.'),
        ('sub', '9.4  Cancellation & Refunds',
         'You may cancel your Subscription at any time by contacting legal@ocypheris.com. Cancellation takes effect at the end of your current billing period. Fees paid for the current period are non-refundable except where Ocypheris has materially failed to deliver the subscribed Services. Pre-paid fees for future periods will be refunded on a pro-rata basis upon written request.'),
    ])

    add_section(doc, 10, 'Confidentiality', [
        'Each party agrees to hold in strict confidence all non-public information disclosed by the other party in connection with the Services ("Confidential Information"), including security Findings, technical architecture details, pricing, and business information. Neither party shall disclose Confidential Information to any third party without the other\'s prior written consent, except as required by law.',
        'EXCEPTIONS: This obligation does not apply to information that: (a) was publicly known at the time of disclosure; (b) becomes publicly known through no breach of these Terms; (c) was independently developed without use of Confidential Information; or (d) is required to be disclosed by law, provided the disclosing party gives prompt written notice where practicable.',
        'This confidentiality obligation survives termination of the Subscription for a period of three (3) years.',
    ])

    add_section(doc, 11, 'Disclaimer of Warranties', [
        'TO THE MAXIMUM EXTENT PERMITTED BY APPLICABLE LAW, THE SERVICES ARE PROVIDED "AS IS" AND "AS AVAILABLE." OCYPHERIS EXPRESSLY DISCLAIMS ALL WARRANTIES, EXPRESS OR IMPLIED, INCLUDING WARRANTIES OF MERCHANTABILITY, FITNESS FOR A PARTICULAR PURPOSE, TITLE, AND NON-INFRINGEMENT.',
        'IN PARTICULAR, OCYPHERIS DOES NOT WARRANT THAT: (A) THE SERVICES WILL IDENTIFY ALL SECURITY VULNERABILITIES OR MISCONFIGURATIONS IN YOUR AWS ENVIRONMENT; (B) IMPLEMENTING ANY REMEDIATION ACTION WILL PREVENT ALL SECURITY INCIDENTS; (C) FINDINGS OR RECOMMENDATIONS ARE FREE FROM ERROR; OR (D) THE SERVICES WILL MEET YOUR SPECIFIC REGULATORY OR COMPLIANCE REQUIREMENTS WITHOUT INDEPENDENT ASSESSMENT.',
        'The Services are a decision-support and execution tool. The Customer retains sole responsibility for final security and compliance decisions.',
    ])

    add_section(doc, 12, 'Limitation of Liability', [
        'TO THE MAXIMUM EXTENT PERMITTED BY APPLICABLE LAW, OCYPHERIS SHALL NOT BE LIABLE FOR ANY INDIRECT, INCIDENTAL, SPECIAL, CONSEQUENTIAL, OR PUNITIVE DAMAGES, OR ANY LOSS OF PROFITS, REVENUE, DATA, BUSINESS, OR GOODWILL, ARISING OUT OF OR IN CONNECTION WITH THE SERVICES OR THESE TERMS, EVEN IF ADVISED OF THE POSSIBILITY OF SUCH DAMAGES.',
        'OCYPHERIS\'S TOTAL AGGREGATE LIABILITY TO YOU FOR ALL CLAIMS ARISING OUT OF OR RELATED TO THE SERVICES SHALL NOT EXCEED THE TOTAL FEES ACTUALLY PAID BY YOU TO OCYPHERIS DURING THE THREE (3) MONTHS PRECEDING THE EVENT GIVING RISE TO THE CLAIM.',
        'THIS LIMITATION APPLIES TO ALL CLAIMS, WHETHER BASED IN CONTRACT, TORT (INCLUDING NEGLIGENCE), STRICT LIABILITY, OR ANY OTHER LEGAL THEORY.',
    ])

    add_section(doc, 13, 'Indemnification', [
        'BY CUSTOMER: You agree to indemnify, defend, and hold harmless Ocypheris and its officers, employees, directors, and contractors from and against any claims, damages, losses, costs, and expenses (including reasonable legal fees) arising out of or related to: (a) your use of the Services in violation of these Terms; (b) your approval and initiation of any Remediation Run; (c) your breach of any representation, warranty, or obligation under these Terms; or (d) your violation of applicable law or third-party rights.',
        'BY OCYPHERIS: Ocypheris agrees to indemnify, defend, and hold you harmless from any third-party claims alleging that the Autopilot platform infringes that third party\'s intellectual property rights, provided that you promptly notify us, grant us sole control of the defence, and cooperate reasonably. This does not apply to claims arising from your modifications or combinations with other products.',
    ])

    add_section(doc, 14, 'Term & Termination', [
        ('sub', '14.1  Term',
         'These Terms are effective from the date you first access the Services and remain in force for the duration of your active Subscription and any renewal periods.'),
        ('sub', '14.2  Termination for Cause',
         'Either party may terminate these Terms upon material breach by the other party, if such breach is not cured within 14 days of written notice specifying the breach. Either party may terminate immediately if the other becomes insolvent or subject to bankruptcy proceedings.'),
        ('sub', '14.3  Termination by Ocypheris',
         'Ocypheris may suspend or terminate your access without prior notice if we reasonably determine you are in material violation of Section 7 (Acceptable Use), or if required to do so by applicable law.'),
        ('sub', '14.4  Effect of Termination',
         'Upon termination or expiry:',
         ['Your access to the platform will be disabled',
          'You must immediately revoke all IAM roles granted to Ocypheris in your AWS account',
          'You will have 30 days to request an export of your data; thereafter, we will delete your data per our Privacy Policy',
          'All outstanding fees become immediately due and payable',
          'Sections 6, 8, 10, 11, 12, 13, 14.4, and 15 survive termination']),
    ])

    add_section(doc, 15, 'Governing Law & Disputes', [
        'These Terms are governed by and construed in accordance with the laws of the Arab Republic of Egypt, without regard to its conflict of law principles. Any dispute, controversy, or claim arising out of or relating to these Terms or the Services that cannot be resolved by good-faith negotiation within 30 days shall be finally settled by binding arbitration conducted in Cairo, Egypt, in the English language, under the rules of the Cairo Regional Centre for International Commercial Arbitration (CRCICA). Judgment on the arbitration award may be entered in any court of competent jurisdiction.',
        'Nothing in this section prevents either party from seeking emergency injunctive or other equitable relief from a court of competent jurisdiction to protect intellectual property rights or prevent imminent irreparable harm pending arbitration.',
    ])

    add_section(doc, 16, 'Changes to These Terms', [
        'Ocypheris reserves the right to update these Terms at any time. We will notify you of material changes by email to your registered address and by posting a notice on ocypheris.com at least 14 days before the updated Terms take effect. Your continued use of the Services after the effective date constitutes your acceptance of the updated Terms. If you do not agree, you must stop using the Services and notify us to close your account.',
    ])

    add_section(doc, 17, 'Miscellaneous', [
        ['Entire Agreement: These Terms, together with the Privacy Policy and any applicable order form, constitute the entire agreement between the parties and supersede all prior negotiations, representations, or agreements.',
         'Severability: If any provision is held invalid or unenforceable, it will be modified to the minimum extent necessary; remaining provisions remain in full force.',
         'Waiver: Failure to enforce any right or provision does not constitute a waiver of that right or provision.',
         'Assignment: You may not assign these Terms without our prior written consent. Ocypheris may assign these Terms in connection with a merger, acquisition, or asset sale with notice to you.',
         'Force Majeure: Neither party shall be liable for failure to perform due to causes beyond their reasonable control, including AWS regional outages, natural disasters, acts of government, or cyber-attacks.',
         'Notices: All legal notices must be sent to legal@ocypheris.com (for Ocypheris) or to the account email registered in your Autopilot profile (for you).'],
    ])

    add_section(doc, 18, 'Contact', [
        'For questions regarding these Terms or your account, contact us at:',
        'Ocypheris — Legal & Compliance\nEmail: legal@ocypheris.com\nWebsite: ocypheris.com',
    ])

    out_path = os.path.join(OUTPUT_DIR, 'Ocypheris_Terms_and_Conditions.docx')
    doc.save(out_path)
    print(f'Saved: {out_path}')


if __name__ == '__main__':
    build_privacy_policy()
    build_terms()
    print('Done. Both documents written to docs/legal/')
