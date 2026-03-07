import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '4F46E5')
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    new_run.append(rPr)
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def create_cookie_policy():
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    heading = doc.add_heading('Cookie Policy', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Effective Date: December 13, 2025').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Ocypheris Autopilot').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    def add_section(doc, num, title, paragraphs):
        h = doc.add_heading(f"{num}. {title}", level=1)
        for p in paragraphs:
            if isinstance(p, list):
                for item in p:
                    doc.add_paragraph(item, style='List Bullet')
            else:
                doc.add_paragraph(p)

    add_section(doc, 1, 'Introduction', [
        'This Cookie Policy explains how Ocypheris ("Ocypheris," "we," "us," or "our") uses cookies and similar storage technologies when you visit or use our website at ocypheris.com and the Autopilot platform (the "Services").',
        'By using our Services, you agree to the use of these technologies as described in this policy. This statement should be read alongside our Privacy Policy, which explains how we protect your personal data entirely.'
    ])

    add_section(doc, 2, 'What Are Cookies and Local Storage?', [
        'Cookies are small text files placed on your device by websites you visit. They are widely used to make websites work, improve efficiency, and provide information to the site owners.',
        'Local Storage (Web Storage) is an industry-standard technology that allows a website or application to store data locally on your computer or mobile device. Unlike cookies, local storage data is not automatically transmitted to the server with every HTTP request, making it more secure and efficient for storing application state.'
    ])

    add_section(doc, 3, 'How We Use These Technologies', [
        'Ocypheris Autopilot is designed as a secure, privacy-first B2B platform. We explicitly use minimal tracking. We use cookies and local storage exclusively for essential platform functionality, security, and user preference retention.',
        'We do not use third-party advertising cookies, cross-site trackers, or data-broker pixels.'
    ])

    add_section(doc, 4, 'Strictly Necessary Technologies', [
        'These technologies are absolutely essential for the Autopilot platform to function. Because they are strictly necessary, they do not require user consent under GDPR/ePrivacy regulations. You cannot opt out of these while using the authenticated platform.',
        [
            "Authentication Tokens (Local Storage): When you log in, we store a secure JSON Web Token (JWT) in your browser's local storage. This verifies your identity to our API and keeps you logged in across page reloads.",
            "CSRF Tokens (Cookies): Short-lived security cookies used to prevent Cross-Site Request Forgery attacks, ensuring that API requests genuinely originate from your browser.",
            "Session Routing (Cookies): Infrastructure cookies deployed by our cloud providers (such as Cloudflare or AWS) to ensure your traffic is securely and reliably routed to the correct server instances."
        ]
    ])

    add_section(doc, 5, 'Functional Preferences', [
        'We use local storage to remember your UI choices so you do not have to reconfigure them on every visit. These are strictly functional and do not track your behaviour across the internet.',
        [
            'Theme Preference (Local Storage): Remembers whether you have selected "light" or "dark" mode.',
            'UI State (Local Storage): Remembers the collapsed/expanded state of your sidebar navigation and data table sorting preferences.'
        ]
    ])

    add_section(doc, 6, 'Analytics & Performance (Currently Disabled)', [
        'At this time, Ocypheris does not deploy third-party analytics cookies (e.g., Google Analytics). We rely exclusively on server-side aggregated telemetry (which does not use cookies) to monitor platform health and performance.',
        'If we introduce first-party or third-party analytics cookies in the future to understand how users interact with the marketing site, we will update this policy and present a consent banner requiring your explicit "opt-in" before such cookies are placed, in accordance with EU/UK law.'
    ])

    add_section(doc, 7, 'Managing Your Settings', [
        'You have full control over the data stored in your browser. You can view, manage, and delete cookies and local storage at any time through your browser settings.',
        'If you choose to block or clear local storage while using the Autopilot platform:',
        [
            'You will be immediately logged out of your active session.',
            'Your UI preferences (like dark mode) will be reset to defaults.',
            'You will need to log in again to restore access.'
        ],
        'To clear your local storage specifically for Autopilot, you can click the "Log Out" button within the platform, which programmatically purges your authentication tokens and session data.'
    ])

    add_section(doc, 8, 'Changes to This Policy', [
        'We may update this Cookie Policy periodically to reflect technological changes, new platform features, or regulatory requirements. We will notify you of any material changes by posting the updated policy on our website with a new "Effective Date."'
    ])

    add_section(doc, 9, 'Contact Us', [
        'If you have any questions about this Cookie Policy or our data practices, please contact our privacy team:',
        'Ocypheris — Data Privacy Team\nEmail: legal@ocypheris.com\nWebsite: ocypheris.com'
    ])

    outpath = os.path.join(os.path.dirname(__file__), '..', 'docs', 'legal', 'Ocypheris_Cookie_Policy.docx')
    doc.save(outpath)
    print(f"Saved: {os.path.abspath(outpath)}")

if __name__ == '__main__':
    create_cookie_policy()
